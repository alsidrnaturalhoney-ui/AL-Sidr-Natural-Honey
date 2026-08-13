from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

source = Path('/home/ubuntu/al-sidr-brand-os/knowledge-base/brand-core.md').read_text()
doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Arial'
styles['Normal'].font.size = Pt(10)

title = doc.add_heading('Al Sidr Natural Honey — Brand Core Knowledge Base', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Canonical internal reference assembled from the Al Sidr brand operating system.')

for line in source.splitlines():
    if not line.strip():
        continue
    if line.startswith('# '):
        continue
    if line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=2)
    elif line.startswith('|') or line.startswith('---'):
        continue
    elif line.startswith('> '):
        p = doc.add_paragraph()
        p.style = doc.styles['Intense Quote']
        p.add_run(line[2:].strip())
    elif line.startswith('- '):
        doc.add_paragraph(line[2:].strip(), style='List Bullet')
    else:
        doc.add_paragraph(line.strip())

out = '/home/ubuntu/al-sidr-brand-os/Al-Sidr-Brand-Core-Knowledge-Base.docx'
doc.save(out)
print(out)
